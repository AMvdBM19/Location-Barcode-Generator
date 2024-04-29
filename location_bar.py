from io import BytesIO

from barcode import Code128

from barcode.writer import ImageWriter

from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx import Document

from docx.shared import Pt

from docx.enum.section import WD_ORIENT

from docx.shared import Inches

from docx.shared import Mm

from docx.shared import Cm

import time

import easygui

import os

from colorama import Fore, Back, Style

from colorama import just_fix_windows_console

from docx.enum.text import WD_ALIGN_PARAGRAPH

from barcode.writer import ImageWriter

from PIL import Image as PIL_Image

from PIL import ImageFont

from barcode.writer import ImageWriter

just_fix_windows_console()

document = Document()

section = document.sections


print("")

print(Fore.GREEN + "Welcome to GPC Location-Barcode Label Generator")

time.sleep(1)

print(Fore.WHITE + "")

print("What format of labels do you wish to create?")
print("")
print("Please choose between UPS size labels (4inch x 6inch) and 1.4inch x 1.6inch labels.")
print("")
time.sleep(1)

#label_format input parameters
medium =("medium", "m", "middel")
large = ("large", "l", "groot", "g")


while True:

    print("Type ''large'' for 4inch x 6inch labels; ''medium'' for 1.4inch x 1.6inch labels ")
    print("or ''exit'' if you want to close the program.")
    print("")

    label_format = input(":").lower()
    print("")
    label_format = str(label_format)


    if label_format in large:

        section = document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE

        section.page_height = Mm(101.6)
        section.page_width = Mm(152)

        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(00)
        section.bottom_margin = Mm(0)

        break

    elif label_format in medium:

        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT

        section.page_height = Mm(35.0)
        section.page_width = Mm(37.5)

        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)

        break

    elif label_format == "exit" or label_format == "quit":

        time.sleep(2)

        break

    else:

        print(Fore.RED + "User input not valid")
        print(Fore.WHITE + "")
        time.sleep(2)

        continue


location_id = ('1') #This value can and should be changed if necessary!
print("")
zone = input("Please input the Zone (01, 02, etc): ")
print("")
rack = input("Please input the Rack (A, B, C, etc): ").upper()

print("")
row = input("Please input the Row (1, 2, 3, etc): ")
print("")
locations = input("Please input the amount of locations in the row: ")
locations = int(locations)
print("")


medium_page_height = Cm(3.50)
medium_page_width = Cm(3.75)

for result in range(1, locations + 1):

    if label_format in medium:

        loc_code = f"{location_id}-{zone}-{rack}-{row}-{result}"

        loc_bar = Code128(loc_code, writer=ImageWriter())

        barcode_bytes = BytesIO()
        loc_bar.write(barcode_bytes)
        barcode_bytes.seek(0)

        print(Fore.YELLOW + loc_code)


        with PIL_Image.open(barcode_bytes) as img:
            img = img.resize((int(medium_page_width.inches * 70), int(medium_page_height.inches * 70)))
            resized_barcode_bytes = BytesIO()
            img.save(resized_barcode_bytes, format='PNG')
            resized_barcode_bytes.seek(0)

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()

        r.add_picture(resized_barcode_bytes)

    else:

        loc_code = f"{location_id}-{zone}-{rack}-{row}-{result}"

        loc_bar = Code128(loc_code, writer=ImageWriter())

        barcode_bytes = BytesIO()
        loc_bar.write(barcode_bytes)
        barcode_bytes.seek(0)

        print(Fore.YELLOW + loc_code)

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()


        r.add_picture(barcode_bytes)

time.sleep(1)
document.save(f"{zone}-{rack}{row}.docx")

#By Andres van den Bos 2024
